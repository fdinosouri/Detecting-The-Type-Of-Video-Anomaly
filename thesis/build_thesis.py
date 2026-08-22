# -*- coding: utf-8 -*-
"""Assemble the thesis into a Word document.

    python thesis/make_figures.py     # once, to draw the plots
    python thesis/build_thesis.py

The layout follows the Hamedan University of Technology writing guide:
front matter without page numbers, contents pages numbered with abjad
letters, the main text numbered with digits, and the appendices, English
abstract and English title page unnumbered at the end. Each chapter is
its own Word section so it can carry its own running header and start on
an unnumbered title page.

Fields (the three tables of contents) are inserted as real Word fields.
Word fills them in on open, or with Ctrl+A then F9.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

import content_fa as C
from docx_style import (
    EN, FA, SIZE_BODY, SIZE_CAPTION, SIZE_TABLE,
    _el, _set, add_field, add_page_break, add_section, build_styles,
    configure_section, ltr_paragraph, rtl_paragraph, set_footer_page_number,
    set_header_text, set_run_font,
)

HERE = Path(__file__).resolve().parent
FIGURES = HERE / "figures"
OUTPUT = HERE / "Thesis_HUT_VideoAnomaly.docx"

PERSIAN_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
ORDINALS = ["اول", "دوم", "سوم", "چهارم", "پنجم", "ششم", "هفتم", "هشتم"]

# Latin fragments inside Persian text get reordered by the bidi algorithm:
# "[0.339, 0.503]" comes out with its brackets mirrored, and "0.44 / 0.86"
# comes out reversed. Wrapping such a fragment in an LTR embedding keeps it
# as written.
LRE, POP_DIR = "\u202A", "\u202C"
LATIN_RUN = re.compile(
    r"[\[\(]?[A-Za-z0-9][A-Za-z0-9 \.,:;/\\\-\+\*=_%'\"\(\)\[\]]*"
    r"[A-Za-z0-9\)\]%]"
)


def fa_num(value):
    return str(value).translate(PERSIAN_DIGITS)


def bidi_safe(text):
    """Pin multi-token Latin fragments so bidi does not reorder them."""
    def wrap(match):
        run = match.group(0)

        if any(ch in run for ch in "()[]/ "):
            return f"{LRE}{run}{POP_DIR}"

        return run

    return LATIN_RUN.sub(wrap, text)


class Builder:
    def __init__(self):
        self.doc = Document()
        build_styles(self.doc)
        configure_section(self.doc.sections[0])
        self.doc.sections[0].different_first_page_header_footer = False

        self.prefix = ""          # chapter number or appendix letter
        self.section_no = 0
        self.sub_no = 0
        self.subsub_no = 0
        self.table_no = 0
        self.figure_no = 0
        self.eq_no = 0
        self.after_heading = False

    # ------------------------------------------------------------------
    # primitives
    # ------------------------------------------------------------------

    def para(self, text, style="Normal", *, indent=False, rtl=True,
             sizes=SIZE_BODY, bold=False, align=None, outline=None):
        paragraph = self.doc.add_paragraph(style=style)

        if align is not None:
            paragraph.alignment = align

        if outline is not None:
            # 9 means body text, which keeps a paragraph out of the
            # table of contents even when it carries a heading style
            _set(paragraph.paragraph_format.element.get_or_add_pPr(),
                 "w:outlineLvl", val=outline)

        if rtl:
            rtl_paragraph(paragraph)
        else:
            ltr_paragraph(paragraph)

        if indent:
            paragraph.paragraph_format.first_line_indent = Cm(0.5)

        for chunk in text.split("\n"):
            if paragraph.runs:
                paragraph.add_run().add_break()

            run = paragraph.add_run(bidi_safe(chunk) if rtl else chunk)
            set_run_font(run, sizes, bold=bold, rtl=rtl)

        return paragraph

    def blank(self, count=1):
        for _ in range(count):
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

    def heading(self, level, text):
        number = {
            2: lambda: f"{self.prefix}-{fa_num(self.section_no)}",
            3: lambda: (f"{self.prefix}-{fa_num(self.section_no)}"
                        f"-{fa_num(self.sub_no)}"),
            4: lambda: (f"{self.prefix}-{fa_num(self.section_no)}"
                        f"-{fa_num(self.sub_no)}-{fa_num(self.subsub_no)}"),
        }[level]()

        paragraph = self.doc.add_paragraph(style=f"Heading {level}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rtl_paragraph(paragraph)

        sizes = {2: (16, 14), 3: (14, 13), 4: (13, 12)}[level]
        run = paragraph.add_run(bidi_safe(f"{number} {text}"))
        set_run_font(run, sizes, bold=True)

        self.after_heading = True

    # ------------------------------------------------------------------
    # tables, figures, equations
    # ------------------------------------------------------------------

    def table(self, caption, headers, rows, *, latin_body=False):
        self.table_no += 1
        label = f"جدول {self.prefix}-{fa_num(self.table_no)}  {caption}"

        cap = self.doc.add_paragraph(style="Table Title*")
        rtl_paragraph(cap)
        run = cap.add_run(bidi_safe(label))
        set_run_font(run, SIZE_CAPTION, bold=True)

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # right-to-left column order
        _set(table._tbl.tblPr, "w:bidiVisual")

        for cell, text in zip(table.rows[0].cells, headers):
            self._cell(cell, text, bold=True, rtl=True)

        # a table that spans a page break repeats its header there
        header_pr = table.rows[0]._tr.get_or_add_trPr()
        _set(header_pr, "w:tblHeader")

        for row in rows:
            new_row = table.add_row()
            _set(new_row._tr.get_or_add_trPr(), "w:cantSplit")
            cells = new_row.cells

            for index, (cell, text) in enumerate(zip(cells, row)):
                rtl = not (latin_body and index == len(row) - 1)
                self._cell(cell, str(text), rtl=rtl)

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(18)
        spacer.paragraph_format.line_spacing = 1.0
        self.after_heading = False

    def _cell(self, cell, text, *, bold=False, rtl=True):
        paragraph = cell.paragraphs[0]
        paragraph.style = self.doc.styles["Thesis Table Text"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if rtl:
            rtl_paragraph(paragraph)
        else:
            ltr_paragraph(paragraph)

        run = paragraph.add_run(bidi_safe(text) if rtl else text)
        set_run_font(run, SIZE_TABLE, bold=bold, rtl=rtl)

    def figure(self, filename, caption):
        path = FIGURES / filename

        if not path.exists():
            raise SystemExit(
                f"missing figure {path}. Run: python thesis/make_figures.py"
            )

        self.figure_no += 1

        holder = self.doc.add_paragraph(style="Thesis Figure")
        holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        holder.add_run().add_picture(str(path), width=Cm(14.0))

        cap = self.doc.add_paragraph(style="Pic Title*")
        rtl_paragraph(cap)
        run = cap.add_run(bidi_safe(
            f"شکل {self.prefix}-{fa_num(self.figure_no)}  {caption}"
        ))
        set_run_font(run, SIZE_CAPTION, bold=True)
        self.after_heading = False

    def equation(self, text):
        self.eq_no += 1

        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set(table._tbl.tblPr, "w:bidiVisual")

        # no borders -- the table is only here to pin the number left
        borders = _set(table._tbl.tblPr, "w:tblBorders")

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            borders.append(_el(f"w:{edge}", val="none", sz="0", space="0"))

        table.columns[0].width = Cm(12.5)
        table.columns[1].width = Cm(2.5)

        body = table.rows[0].cells[0].paragraphs[0]
        body.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ltr_paragraph(body)
        run = body.add_run(text)
        set_run_font(run, SIZE_BODY, rtl=False)
        run.font.italic = True

        number = table.rows[0].cells[1].paragraphs[0]
        number.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rtl_paragraph(number)
        run = number.add_run(f"({self.prefix}-{fa_num(self.eq_no)})")
        set_run_font(run, SIZE_BODY)

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.line_spacing = 1.0
        self.after_heading = False

    def code(self, text):
        for line in text.split("\n"):
            paragraph = self.doc.add_paragraph(style="Thesis Code")
            ltr_paragraph(paragraph)
            run = paragraph.add_run(line if line else " ")
            set_run_font(run, (11, 10), latin="Courier New",
                         persian="Courier New", rtl=False)

        self.after_heading = False

    def bullets(self, items, numbered=False):
        for index, item in enumerate(items, start=1):
            marker = f"{fa_num(index)}. " if numbered else "– "
            paragraph = self.para(marker + item, style="Thesis Bullet")
            paragraph.paragraph_format.right_indent = Cm(0.9)

        self.after_heading = False

    def quote(self, text):
        paragraph = self.para(text)
        paragraph.paragraph_format.right_indent = Cm(1.0)
        paragraph.paragraph_format.left_indent = Cm(1.0)

        for run in paragraph.runs:
            run.font.italic = True

        self.after_heading = False

    # ------------------------------------------------------------------
    # block dispatch
    # ------------------------------------------------------------------

    def render(self, blocks):
        for block in blocks:
            kind = block[0]

            if kind == "h2":
                self.section_no += 1
                self.sub_no = 0
                self.subsub_no = 0
                self.heading(2, block[1])

            elif kind == "h3":
                self.sub_no += 1
                self.subsub_no = 0
                self.heading(3, block[1])

            elif kind == "h4":
                self.subsub_no += 1
                self.heading(4, block[1])

            elif kind == "p":
                self.para(block[1], indent=self.after_heading)
                self.after_heading = False

            elif kind == "bul":
                self.bullets(block[1])

            elif kind == "num":
                self.bullets(block[1], numbered=True)

            elif kind in ("tbl", "tbl_a"):
                self.table(block[1], block[2], block[3],
                           latin_body=(kind == "tbl_a"))

            elif kind == "fig":
                self.figure(block[1], block[2])

            elif kind == "eq":
                self.equation(block[1])

            elif kind == "code":
                self.code(block[1])

            elif kind == "quote":
                self.quote(block[1])

            else:
                raise ValueError(f"unknown block kind: {kind}")

    # ------------------------------------------------------------------
    # front matter
    # ------------------------------------------------------------------

    def bismillah(self):
        self.blank(8)
        self.para("بِسْمِ اللهِ الرَّحْمنِ الرَّحیمِ", style="Thesis Title Line",
                  sizes=(22, 20), bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    def title_page_fa(self):
        add_page_break(self.doc)
        meta = C.META

        self.para(meta["university"], style="Thesis Title Line", sizes=(16, 15),
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.para(meta["department"], style="Thesis Title Line", sizes=(14, 13),
                  bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.blank(1)
        self.para(
            f"پایان‌نامه ارائه شده به عنوان بخشی از ملزومات، جهت دریافت درجه‌ی "
            f"{meta['degree']} در رشته‌ی {meta['field']} – گرایش "
            f"{meta['orientation']}",
            sizes=(13, 12), align=WD_ALIGN_PARAGRAPH.CENTER)
        self.blank(1)
        self.para(meta["title"], style="Thesis Title Line", sizes=(18, 16),
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.blank(1)
        self.para("توسط:", sizes=(14, 13), align=WD_ALIGN_PARAGRAPH.CENTER)
        self.para(meta["author"], style="Thesis Title Line", sizes=(16, 15),
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.blank(2)
        self.para("استاد راهنما:", sizes=(14, 13),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        self.para(meta["supervisor"], style="Thesis Title Line", sizes=(16, 15),
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.blank(2)
        self.para(meta["date"], sizes=(14, 13),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    def copyright_page(self):
        add_page_break(self.doc)
        self.blank(10)
        self.para(
            "کلیه‌ی حقوق مادی مترتب بر نتایج مطالعات، ابتکارات و نوآوری‌های ناشی "
            "از تحقیق موضوع این پایان‌نامه متعلق به دانشگاه صنعتی همدان است.",
            style="Thesis Title Line", sizes=(14, 13), bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER)

    def dedication(self):
        add_page_break(self.doc)
        self.blank(9)

        for line in C.DEDICATION:
            self.para(line, style="Thesis Title Line", sizes=(15, 14),
                      bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    def acknowledgement(self):
        add_page_break(self.doc)
        self.blank(2)
        self.para("تقدیر و تشکر", style="Thesis Title Line", sizes=(18, 16),
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.blank(1)

        for paragraph in C.ACKNOWLEDGEMENT:
            self.para(paragraph, indent=True)

    def abstract_fa(self):
        add_page_break(self.doc)
        self.para("چکیده", style="Thesis Title Line", sizes=(18, 16), bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        for paragraph in C.ABSTRACT_FA:
            self.para(paragraph, style="Thesis Abstract", indent=True)

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
        spacer.paragraph_format.line_spacing = 1.0

        paragraph = self.doc.add_paragraph(style="Thesis Abstract")
        rtl_paragraph(paragraph)
        label = paragraph.add_run("واژگان کلیدی: ")
        set_run_font(label, SIZE_BODY, bold=True)
        body = paragraph.add_run(C.KEYWORDS_FA)
        set_run_font(body, SIZE_BODY)

    # ------------------------------------------------------------------
    # contents
    # ------------------------------------------------------------------

    def contents(self):
        section = add_section(self.doc, page_format="arabicAbjad", page_start=1)
        set_header_text(section, "")
        set_footer_page_number(section, enabled=True)

        entries = [
            ("فهرست مطالب", 'TOC \\o "1-4" \\h \\z \\u'),
            ("فهرست جداول", 'TOC \\h \\z \\t "Table Title*,1"'),
            ("فهرست شکل‌ها", 'TOC \\h \\z \\t "Pic Title*,1"'),
        ]

        for index, (title, instruction) in enumerate(entries):
            if index:
                add_page_break(self.doc)

            self.para(title, style="Thesis Title Line", sizes=(18, 16),
                      bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            self.blank(1)

            paragraph = self.doc.add_paragraph()
            rtl_paragraph(paragraph)
            add_field(
                paragraph, instruction,
                "برای نمایش این فهرست، در Word کلیدهای Ctrl+A و سپس F9 را بزنید.",
            )

        add_page_break(self.doc)
        self.para("فهرست علائم و اختصارات", style="Thesis Title Line",
                  sizes=(18, 16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.blank(1)

        self.prefix = ""
        table = self.doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set(table._tbl.tblPr, "w:bidiVisual")

        for cell, text in zip(table.rows[0].cells,
                              ["نماد / اختصار", "عبارت کامل", "معادل فارسی"]):
            self._cell(cell, text, bold=True)

        symbols = [
            ("VAD", "Video Anomaly Detection", "تشخیص ناهنجاری در ویدیو"),
            ("MIL", "Multiple Instance Learning", "یادگیری چندنمونه‌ای"),
            ("AUC", "Area Under the ROC Curve", "سطح زیر منحنی مشخصه‌ی عملکرد"),
            ("ViT", "Vision Transformer", "ترنسفورمر بینایی"),
            ("MAE", "Masked Autoencoder", "خودرمزگذار نقاب‌دار"),
            ("CE", "Cross Entropy", "آنتروپی متقاطع"),
            ("CI", "Confidence Interval", "بازه‌ی اطمینان"),
            ("T", "number of clips per video", "تعداد کلیپ در هر ویدیو"),
            ("k", "top-k pooling size", "تعداد کلیپ‌های برتر در تجمیع"),
            ("π_c", "class prior", "احتمال پیشین کلاس"),
            ("w_c", "class weight", "وزن کلاس"),
            ("α", "mixup / logit-adjust coefficient",
             "ضریب آمیزش یا تنظیم لاجیت"),
            ("λ", "mixup interpolation factor", "ضریب ترکیب آمیزش"),
            ("τ", "logit scale", "مقیاس لاجیت"),
        ]

        for symbol, full, persian in symbols:
            cells = table.add_row().cells
            self._cell(cells[0], symbol, rtl=False)
            self._cell(cells[1], full, rtl=False)
            self._cell(cells[2], persian)

    # ------------------------------------------------------------------
    # chapters
    # ------------------------------------------------------------------

    def chapter(self, index, chapter, first=False):
        title = chapter["title"]
        header = f"فصل {ORDINALS[index - 1]}: {title}"

        section = add_section(
            self.doc,
            page_format="decimal" if first else None,
            page_start=1 if first else None,
            different_first_page=True,
        )
        set_header_text(section, header)
        set_footer_page_number(section, enabled=True)

        # chapter title page -- no header, no page number
        self.blank(7)
        self.para(f"فصل {ORDINALS[index - 1]}", style="Heading 1",
                  sizes=(28, 15), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  outline=9)
        self.para(title, style="Heading 1", sizes=(28, 15), bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_page_break(self.doc)

        self.prefix = fa_num(index)
        self.section_no = 0
        self.sub_no = 0
        self.subsub_no = 0
        self.table_no = 0
        self.figure_no = 0
        self.eq_no = 0
        self.after_heading = False

        self.render(chapter["blocks"])

    # ------------------------------------------------------------------
    # back matter
    # ------------------------------------------------------------------

    def references(self):
        section = add_section(self.doc, different_first_page=False)
        set_header_text(section, "منابع و مآخذ")
        set_footer_page_number(section, enabled=True)

        self.para("منابع و مآخذ", style="Thesis Title Line", sizes=(20, 18),
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, outline=0)
        self.blank(1)

        for index, reference in enumerate(C.REFERENCES, start=1):
            paragraph = self.doc.add_paragraph(style="Thesis Reference EN")
            ltr_paragraph(paragraph)
            paragraph.paragraph_format.left_indent = Cm(0.9)
            paragraph.paragraph_format.first_line_indent = Cm(-0.9)

            run = paragraph.add_run(f"[{index}] {reference}")
            set_run_font(run, (13, 12), rtl=False)

    def appendices(self):
        section = add_section(self.doc, different_first_page=True)
        set_header_text(section, "پیوست‌ها")
        set_footer_page_number(section, enabled=False)

        self.blank(8)
        self.para("پیوست‌ها", style="Heading 1", sizes=(28, 15), bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        for appendix in C.APPENDICES:
            add_page_break(self.doc)

            self.prefix = appendix["letter"]
            self.section_no = 0
            self.sub_no = 0
            self.subsub_no = 0
            self.table_no = 0
            self.figure_no = 0
            self.eq_no = 0
            self.after_heading = False

            self.para(f"پیوست {appendix['letter']}: {appendix['title']}",
                      style="Thesis Title Line", sizes=(18, 16), bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, outline=1)
            self.blank(1)
            self.render(appendix["blocks"])

    def abstract_en(self):
        section = add_section(self.doc, rtl=False, different_first_page=False)
        set_header_text(section, "")
        set_footer_page_number(section, enabled=False)

        self.para("Abstract", style="Thesis Title Line EN", sizes=(18, 16),
                  bold=True, rtl=False, align=WD_ALIGN_PARAGRAPH.CENTER)

        for paragraph in C.ABSTRACT_EN:
            self.para(paragraph, style="Thesis Abstract EN", rtl=False,
                      indent=True)

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
        spacer.paragraph_format.line_spacing = 1.0

        paragraph = self.doc.add_paragraph(style="Thesis Abstract EN")
        ltr_paragraph(paragraph)
        label = paragraph.add_run("Keywords: ")
        set_run_font(label, SIZE_BODY, bold=True, rtl=False)
        body = paragraph.add_run(C.KEYWORDS_EN)
        set_run_font(body, SIZE_BODY, rtl=False)

    def title_page_en(self):
        section = add_section(self.doc, rtl=False, different_first_page=False)
        set_header_text(section, "")
        set_footer_page_number(section, enabled=False)

        meta = C.META

        def line(text, size=(14, 13), bold=False):
            self.para(text, style="Thesis Title Line EN", sizes=size,
                      bold=bold, rtl=False, align=WD_ALIGN_PARAGRAPH.CENTER)

        line("Hamedan University of Technology", (16, 15), True)
        line(meta["department_en"])
        self.blank(1)
        line("Submitted in Partial Fulfillment of the Requirements for the "
             f"Degree of {meta['degree_en']} in {meta['field_en']}", (13, 12))
        self.blank(1)
        line(meta["title_en"], (18, 16), True)
        self.blank(1)
        line("By")
        line(meta["author_en"], (16, 15), True)
        self.blank(1)
        line("Supervisor")
        line(meta["supervisor_en"], (16, 15), True)
        self.blank(2)
        line(meta["date_en"])

    # ------------------------------------------------------------------

    def build(self):
        set_footer_page_number(self.doc.sections[0], enabled=False)

        self.bismillah()
        self.title_page_fa()
        self.copyright_page()
        self.dedication()
        self.acknowledgement()
        self.abstract_fa()

        self.contents()

        for index, chapter in enumerate(C.CHAPTERS, start=1):
            self.chapter(index, chapter, first=(index == 1))

        self.references()
        self.appendices()
        self.abstract_en()
        self.title_page_en()

        self.doc.save(OUTPUT)

        return OUTPUT


def main():
    path = Builder().build()
    size = path.stat().st_size / 1024
    print(f"wrote {path}  ({size:.0f} KB)")
    print("open it in Word, press Ctrl+A then F9 to fill the three contents "
          "pages, and install the B Nazanin font if it is not present.")


if __name__ == "__main__":
    main()
