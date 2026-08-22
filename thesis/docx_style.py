"""Word styling that follows the Hamedan University of Technology guide.

Everything the guide fixes numerically -- margins, font sizes, line and
paragraph spacing, page-number formats, header placement -- lives here so
the content module can stay plain text.

The two font families are B Nazanin for Persian and Times New Roman for
Latin. Word picks between them per character run using the "complex
script" font slot, so every run sets both: ``w:cs`` carries the Persian
face and size, ``w:ascii``/``w:hAnsi`` the Latin one.
"""

from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

FA = "B Nazanin"
EN = "Times New Roman"
MONO = "Courier New"

# guide table 1-1 / 1-2: (Persian pt, Latin pt)
SIZE_BODY = (14, 13)
SIZE_H1 = (28, 15)
SIZE_H2 = (16, 14)
SIZE_H3 = (14, 13)
SIZE_H4 = (13, 12)
SIZE_CAPTION = (13, 12)
SIZE_TABLE = (13, 11)
SIZE_HEADER = (11, 11)
SIZE_FOOTNOTE = (10, 10)
SIZE_REF = (13, 12)
SIZE_CODE = (11, 10)


# --------------------------------------------------------------------------
# low-level OOXML helpers
# --------------------------------------------------------------------------

def _el(tag, **attrs):
    node = OxmlElement(tag)

    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), str(value))

    return node


# WordprocessingML validates child order, and Word/LibreOffice reject a
# document whose pPr or rPr children are out of sequence -- so a new child
# has to be inserted before its successors rather than appended.
_ORDER = {
    "pPr": (
        "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
        "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
        "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku",
        "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
        "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
        "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
        "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
        "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
        "w:rPr", "w:sectPr", "w:pPrChange",
    ),
    "rPr": (
        "w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
        "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
        "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
        "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern", "w:position",
        "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd",
        "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang",
        "w:eastAsianLayout", "w:specVanish", "w:oMath",
    ),
    "sectPr": (
        "w:headerReference", "w:footerReference", "w:footnotePr",
        "w:endnotePr", "w:type", "w:pgSz", "w:pgMar", "w:paperSrc",
        "w:pgBorders", "w:lnNumType", "w:pgNumType", "w:cols", "w:formProt",
        "w:vAlign", "w:noEndnote", "w:titlePg", "w:textDirection", "w:bidi",
        "w:rtlGutter", "w:docGrid", "w:printerSettings", "w:sectPrChange",
    ),
    "trPr": (
        "w:cnfStyle", "w:divId", "w:gridBefore", "w:gridAfter", "w:wBefore",
        "w:wAfter", "w:cantSplit", "w:trHeight", "w:tblHeader",
        "w:tblCellSpacing", "w:jc", "w:hidden", "w:ins", "w:del",
        "w:trPrChange",
    ),
    "tblPr": (
        "w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
        "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
        "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd",
        "w:tblLayout", "w:tblCellMar", "w:tblLook", "w:tblCaption",
        "w:tblDescription", "w:tblPrChange",
    ),
}


def _local(element):
    return element.tag.split("}")[-1]


def _set(parent, tag, **attrs):
    """Set ``tag`` on ``parent``, creating it in schema order if absent."""
    node = parent.find(qn(tag))

    if node is None:
        node = OxmlElement(tag)
        order = _ORDER.get(_local(parent))

        if order and tag in order:
            successors = order[order.index(tag) + 1:]
            parent.insert_element_before(node, *successors)
        else:
            parent.append(node)

    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), str(value))

    return node


def rtl_paragraph(paragraph):
    """Mark a paragraph right-to-left."""
    _set(paragraph.paragraph_format.element.get_or_add_pPr(), "w:bidi",
         val="1")


def ltr_paragraph(paragraph):
    _set(paragraph.paragraph_format.element.get_or_add_pPr(), "w:bidi", val="0")


def rtl_run(run):
    _set(run._element.get_or_add_rPr(), "w:rtl", val="1")


def set_run_font(run, sizes=SIZE_BODY, bold=False, latin=EN, persian=FA,
                 rtl=True, colour=None):
    """Apply the paired Persian/Latin font of the guide to one run."""
    persian_pt, latin_pt = sizes

    run.font.name = latin
    run.font.size = Pt(latin_pt)
    run.font.bold = bold

    if colour is not None:
        run.font.color.rgb = RGBColor(*colour)

    rpr = run._element.get_or_add_rPr()
    _set(rpr, "w:rFonts", ascii=latin, hAnsi=latin, cs=persian)
    _set(rpr, "w:szCs", val=int(persian_pt * 2))

    if bold:
        _set(rpr, "w:bCs")

    if rtl:
        _set(rpr, "w:rtl", val="1")
    else:
        _set(rpr, "w:rtl", val="0")


def _style_font(style, sizes, bold=False, latin=EN, persian=FA, rtl=True):
    persian_pt, latin_pt = sizes
    style.font.name = latin
    style.font.size = Pt(latin_pt)
    style.font.bold = bold

    rpr = style.element.get_or_add_rPr()
    _set(rpr, "w:rFonts", ascii=latin, hAnsi=latin, cs=persian)
    _set(rpr, "w:szCs", val=int(persian_pt * 2))

    if bold:
        _set(rpr, "w:bCs")

    if rtl:
        _set(rpr, "w:rtl", val="1")
    else:
        _set(rpr, "w:rtl", val="0")


def _style_paragraph(style, *, space_after=6, space_before=0, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, rtl=True, indent=None,
                     keep_with_next=False):
    fmt = style.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)

    if line == 1.0:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        fmt.line_spacing = line

    fmt.alignment = align
    fmt.keep_with_next = keep_with_next

    if indent is not None:
        fmt.first_line_indent = indent

    if rtl:
        _set(style.element.get_or_add_pPr(), "w:bidi", val="1")
    else:
        _set(style.element.get_or_add_pPr(), "w:bidi", val="0")


# --------------------------------------------------------------------------
# styles
# --------------------------------------------------------------------------

def build_styles(doc):
    """Create every paragraph style the thesis uses."""
    styles = doc.styles

    normal = styles["Normal"]
    _style_font(normal, SIZE_BODY)
    _style_paragraph(normal, space_after=6, line=1.5)

    # Heading 1 is the chapter title, which only ever appears on the
    # chapter's own title page.
    h1 = styles["Heading 1"]
    _style_font(h1, SIZE_H1, bold=True)
    _style_paragraph(h1, space_after=12, space_before=0, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER)
    h1.font.color.rgb = RGBColor(0, 0, 0)

    for name, sizes, after in (
        ("Heading 2", SIZE_H2, 12),
        ("Heading 3", SIZE_H3, 10),
        ("Heading 4", SIZE_H4, 6),
        ("Heading 5", SIZE_H4, 6),
    ):
        style = styles[name]
        _style_font(style, sizes, bold=True)
        _style_paragraph(style, space_after=after, space_before=12, line=1.5,
                         align=WD_ALIGN_PARAGRAPH.RIGHT, keep_with_next=True)
        style.font.color.rgb = RGBColor(0, 0, 0)

    def new(name, base="Normal"):
        try:
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            return styles[name]

    # The guide names these two styles explicitly, because the student is
    # told to build the list of tables and list of figures from them.
    table_title = new("Table Title*")
    _style_font(table_title, SIZE_CAPTION, bold=True)
    _style_paragraph(table_title, space_after=6, space_before=18, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)

    pic_title = new("Pic Title*")
    _style_font(pic_title, SIZE_CAPTION, bold=True)
    _style_paragraph(pic_title, space_after=18, space_before=6, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    body_table = new("Thesis Table Text")
    _style_font(body_table, SIZE_TABLE)
    _style_paragraph(body_table, space_after=2, space_before=2, line=1.0,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    figure = new("Thesis Figure")
    _style_paragraph(figure, space_after=6, space_before=18, line=1.0,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    equation = new("Thesis Equation")
    _style_font(equation, SIZE_BODY, latin=EN)
    _style_paragraph(equation, space_after=12, space_before=12, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    code = new("Thesis Code")
    _style_font(code, SIZE_CODE, latin=MONO, persian=MONO, rtl=False)
    _style_paragraph(code, space_after=2, space_before=2, line=1.0,
                     align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False)

    ref = new("Thesis Reference")
    _style_font(ref, SIZE_REF)
    _style_paragraph(ref, space_after=6, space_before=0, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    ref_en = new("Thesis Reference EN")
    _style_font(ref_en, SIZE_REF, rtl=False)
    _style_paragraph(ref_en, space_after=6, space_before=0, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False)

    abstract = new("Thesis Abstract")
    _style_font(abstract, SIZE_BODY)
    _style_paragraph(abstract, space_after=6, space_before=0, line=1.0,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    abstract_en = new("Thesis Abstract EN")
    _style_font(abstract_en, SIZE_BODY, rtl=False)
    _style_paragraph(abstract_en, space_after=6, space_before=0, line=1.0,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, rtl=False)

    title_line = new("Thesis Title Line")
    _style_font(title_line, (16, 15), bold=True)
    _style_paragraph(title_line, space_after=12, space_before=12, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    title_line_en = new("Thesis Title Line EN")
    _style_font(title_line_en, (16, 15), bold=True, rtl=False)
    _style_paragraph(title_line_en, space_after=12, space_before=12, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER, rtl=False)

    header = new("Thesis Header")
    _style_font(header, SIZE_HEADER)
    _style_paragraph(header, space_after=0, space_before=0, line=1.0,
                     align=WD_ALIGN_PARAGRAPH.RIGHT)

    footer = new("Thesis Footer")
    _style_font(footer, SIZE_HEADER)
    _style_paragraph(footer, space_after=0, space_before=0, line=1.0,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    listing = new("Thesis Bullet")
    _style_font(listing, SIZE_BODY)
    _style_paragraph(listing, space_after=6, space_before=0, line=1.5,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    listing.paragraph_format.right_indent = Cm(0.8)

    return doc


# --------------------------------------------------------------------------
# sections, headers, footers
# --------------------------------------------------------------------------

def configure_section(section, *, rtl=True, page_format=None, page_start=None,
                      restart=False):
    """Apply the guide's margins and page-number settings to a section."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    if rtl:
        section.top_margin = Cm(3.5)
        section.right_margin = Cm(3.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
    else:
        section.top_margin = Cm(3.5)
        section.left_margin = Cm(3.5)
        section.bottom_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    section.header_distance = Cm(0)
    section.footer_distance = Cm(1.5)

    sect_pr = section._sectPr
    if rtl:
        _set(sect_pr, "w:bidi", val="1")
    else:
        _set(sect_pr, "w:bidi", val="0")

    if page_format or page_start is not None or restart:
        attrs = {}

        if page_format:
            attrs["fmt"] = page_format

        if page_start is not None:
            attrs["start"] = page_start

        _set(sect_pr, "w:pgNumType", **attrs)


def add_section(doc, *, rtl=True, page_format=None, page_start=None,
                different_first_page=False):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.different_first_page_header_footer = different_first_page
    configure_section(section, rtl=rtl, page_format=page_format,
                      page_start=page_start)

    return section


def _unlink(part):
    """Break the 'same as previous' link so a section can differ."""
    part.is_linked_to_previous = False


def set_header_text(section, text):
    _unlink(section.header)
    _unlink(section.first_page_header)

    header = section.header.paragraphs[0]
    header.style = section.part.document.styles["Thesis Header"]
    header.text = ""
    run = header.add_run(text)
    set_run_font(run, SIZE_HEADER)
    rtl_paragraph(header)

    # the chapter title page carries neither header nor page number
    first = section.first_page_header.paragraphs[0]
    first.text = ""


def add_page_number(paragraph):
    run = paragraph.add_run()
    rpr = run._element.get_or_add_rPr()
    _set(rpr, "w:rFonts", ascii=EN, hAnsi=EN, cs=FA)
    _set(rpr, "w:szCs", val=22)

    begin = _el("w:fldChar", fldCharType="begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = _el("w:fldChar", fldCharType="end")

    run._element.append(begin)
    run._element.append(instr)
    run._element.append(end)


def set_footer_page_number(section, enabled=True):
    _unlink(section.footer)
    _unlink(section.first_page_footer)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.style = section.part.document.styles["Thesis Footer"]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if enabled:
        add_page_number(footer)

    first = section.first_page_footer.paragraphs[0]
    first.text = ""


def add_field(paragraph, instruction, placeholder):
    """Insert a Word field (used for the three tables of contents)."""
    run = paragraph.add_run()
    begin = _el("w:fldChar", fldCharType="begin", dirty="true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = _el("w:fldChar", fldCharType="separate")
    end = _el("w:fldChar", fldCharType="end")

    run._element.append(begin)
    run._element.append(instr)
    run._element.append(separate)

    text = paragraph.add_run(placeholder)
    set_run_font(text, SIZE_BODY)

    closer = paragraph.add_run()
    closer._element.append(end)


def add_page_break(doc):
    from docx.enum.text import WD_BREAK

    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)

    return paragraph
